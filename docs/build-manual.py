#!/usr/bin/env python3
# Build user manual in both .docx and .pdf from a single source of truth.
# Usage: python3 docs/build-manual.py
# Output: docs/Manuale-TireCheckTire.docx + docs/Manuale-TireCheckTire.pdf

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, PageBreak,
    Table, TableStyle, ListFlowable, ListItem, Image,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# ---------------------------------------------------------------------------
# Content (single source of truth — h1 = chapter, h2 = section, p = paragraph,
# bullets = list, table = key/value table, callout = highlighted box).
# ---------------------------------------------------------------------------

DOC_TITLE = "TireCheckTire — Manuale d'uso"
DOC_SUBTITLE = "The Tire Intelligence Suite · v1.0"
DOC_BYLINE = "by PezzaliApp"

CONTENT = [
    ("cover",),

    ("h1", "1. Che cos'è TireCheckTire"),
    ("p",
     "TireCheckTire è la suite professionale per chi lavora con gli pneumatici: "
     "gommisti, autofficine, officine truck, fleet manager, noleggiatori e "
     "distributori B2B. Unisce in un'unica app tre strumenti che prima erano "
     "separati: diagnosi AI dei pneumatici da foto, preventivatore guidato, "
     "storico clienti e officine vicine."),
    ("p",
     "L'app gira nel browser come PWA (Progressive Web App): si installa sulla "
     "Home del telefono come una normale app, funziona offline per le parti "
     "che non richiedono Internet, e non ha bisogno di un'installazione "
     "tradizionale."),
    ("bullets",
     "Diagnosi AI auto e truck/bus/trailer da foto del battistrada",
     "OCR sul fianco: legge automaticamente misura, DOT e marca",
     "Preventivo in 4 step con PDF firmabile e invio WhatsApp",
     "Scanner QR · EAN · DOT con riconoscimento EPREL",
     "Storico unificato (analisi + preventivi) con export CSV/JSON",
     "Officine vicine via OpenStreetMap (gratis, no API key)",
     "AI multi-provider: Gemini, OpenAI, Claude, Mistral, Ollama"),

    ("h1", "2. Primo avvio — onboarding"),
    ("p",
     "Al primo accesso compare un onboarding in 3 step. Puoi sempre saltare e "
     "configurare dopo da Impostazioni."),
    ("h2", "Step 1 — Benvenuto"),
    ("p",
     "Schermata di presentazione. Tocca 'Avanti' per continuare o 'Salta tutto' "
     "per andare direttamente alla Dashboard."),
    ("h2", "Step 2 — Profilo officina"),
    ("p",
     "Inserisci i dati della tua officina. Tutti i campi sono opzionali: "
     "verranno usati nell'intestazione dei PDF (analisi e preventivi)."),
    ("table",
     ["Campo", "Esempio", "Dove appare"],
     ["Nome officina", "Gomme Bianchi", "Intestazione PDF, header app"],
     ["Telefono", "+39 02 12345678", "PDF analisi + preventivo"],
     ["P.IVA", "12345678901", "PDF preventivo"]),
    ("h2", "Step 3 — Provider AI"),
    ("p",
     "Scegli il provider AI che preferisci (default: Google Gemini, gratuita "
     "per uso personale) e incolla la tua API key. Il pulsante 'Testa "
     "connessione' verifica che la chiave funzioni."),
    ("callout",
     "L'app non include chiavi AI. Ogni utente porta la propria, ed è "
     "memorizzata solo nel browser (localStorage). PezzaliApp non riceve "
     "alcun dato."),

    ("h1", "3. Dashboard (Home)"),
    ("p",
     "La schermata principale. In cima vedi un saluto contestuale (Buongiorno, "
     "Buonasera…) e il nome della tua officina."),
    ("h2", "Switch AUTO / TRUCK"),
    ("p",
     "In alto trovi due pulsanti: AUTO (vetture) e TRUCK (camion, bus, trailer). "
     "La selezione è persistente e determina:"),
    ("bullets",
     "Quali campi appaiono nella schermata 'Nuova analisi' (es. asse, "
     "montaggio gemellato, pressione)",
     "Quale prompt AI viene inviato (l'analisi è specializzata per tipo "
     "veicolo)",
     "Quale catalogo di servizi è proposto nel preventivo"),
    ("h2", "Azioni rapide"),
    ("p", "Quattro card di accesso veloce:"),
    ("bullets",
     "📸 Nuova analisi AI → apre la schermata di scansione",
     "📋 Nuovo preventivo → apre il preventivo guidato",
     "🔍 Scanner QR/EAN → apre la fotocamera per leggere un codice",
     "📍 Officine vicine → cerca gommisti e officine nelle vicinanze"),
    ("h2", "KPI"),
    ("p", "Quattro metriche aggiornate in tempo reale:"),
    ("bullets",
     "Analisi oggi: quante diagnosi AI hai fatto oggi",
     "Preventivi mese: preventivi del mese corrente",
     "Urgenze aperte: numero di pneumatici da SOSTITUIRE o con urgenza alta",
     "Ultimo cliente: chi è stato l'ultimo cliente servito"),
    ("h2", "Urgenze e attività recenti"),
    ("p",
     "Sotto le KPI vedi due liste: le urgenze aperte (pneumatici critici da "
     "sostituire) e le ultime 5 attività in ordine cronologico."),

    ("h1", "4. Analisi AI di un pneumatico"),
    ("p",
     "Il cuore dell'app. In meno di 30 secondi ottieni una diagnosi tecnica "
     "completa da una sola foto del battistrada."),
    ("h2", "Flusso completo"),
    ("table",
     ["Passo", "Azione"],
     ["1", "Tocca 'Scan' nella barra in basso (o 'Nuova analisi AI' dalla Dashboard)"],
     ["2", "Verifica la modalità in alto: AUTO o TRUCK"],
     ["3", "Tocca lo slot '📷 Battistrada' → si apre la fotocamera"],
     ["4", "(Opzionale) Tocca '🔎 Fianco' → l'OCR estrae misura/DOT/marca"],
     ["5", "Completa i dati mezzo (targa, posizione, cliente…)"],
     ["6", "Tocca '🤖 Analizza con AI'"],
     ["7", "Attendi 5-15 secondi → vai alla schermata risultati"]),
    ("h2", "Dati pneumatico"),
    ("p", "I campi che puoi compilare manualmente o lasciare riempire dall'OCR:"),
    ("bullets",
     "Misura ETRTO: 205/55 R16 (auto), 315/80 R22.5 (truck)",
     "DOT: codice di 4 cifre (settimana + anno)",
     "Marca: Michelin, Pirelli, Continental, Bridgestone…",
     "Anno produzione (auto) o Pressione bar (truck)",
     "Note: testo libero per l'operatore"),
    ("h2", "Dati mezzo — Auto"),
    ("bullets",
     "Targa",
     "Marca veicolo (Fiat, BMW…)",
     "Posizione: Ant. SX, Ant. DX, Post. SX, Post. DX, Scorta",
     "Cliente"),
    ("h2", "Dati mezzo — Truck"),
    ("bullets",
     "Targa",
     "Tipo mezzo: Trattore, Motrice, Rimorchio, Semirimorchio, Autocarro, Bus…",
     "Asse: Asse 1, Asse 2… Asse 5",
     "Lato/ruota: SX esterna, SX interna, DX esterna, DX interna, Singola",
     "Tipo asse: Sterzante, Trazione, Trailer, Bus, Pusher, Tag, Centrale",
     "Montaggio: Singolo / Gemellato",
     "Marca mezzo, km mezzo, cliente/flotta"),
    ("h2", "Salto AI"),
    ("p",
     "Se l'AI non è configurata o non ti serve la diagnosi, il pulsante "
     "'Salta AI · Inserisci a mano' porta direttamente al preventivo con i "
     "dati pneumatico già pre-popolati."),

    ("h1", "5. Risultato analisi"),
    ("p",
     "Schermata che si apre subito dopo l'analisi. Mostra l'esito complessivo "
     "in un badge grande colorato:"),
    ("bullets",
     "✅ OK (verde): pneumatico in buone condizioni",
     "⚠ ATTENZIONE (giallo): verificare entro breve",
     "🔴 SOSTITUIRE (rosso): cambio immediato"),
    ("h2", "Cosa trovi nella schermata"),
    ("bullets",
     "Foto del battistrada analizzato",
     "Dati rilevati: misura, DOT, anno, marca",
     "Metriche: profondità mm, tipo usura, condizione fianco, urgenza",
     "Valutazione tecnica AI (commento di 2-4 frasi)",
     "Raccomandazioni cliccabili",
     "Canvas per la firma del cliente"),
    ("h2", "Azioni disponibili"),
    ("bullets",
     "📄 Genera Report PDF — scarica un PDF firmabile con dati + foto + firma",
     "📋 Crea preventivo da questa analisi — passa al preventivo con tutto già pre-compilato",
     "📡 Invia webhook Make.com — manda i dati a un webhook configurato",
     "📍 Officine vicine — apre la lista delle officine intorno"),
    ("callout",
     "Per scaricare il PDF devi firmare oppure premere 'Procedi senza firma'. "
     "La firma viene incorporata nel PDF."),

    ("h1", "6. Preventivo"),
    ("p",
     "Il preventivatore è guidato in 4 step. Lo stepper a pallini in alto "
     "ti mostra dove sei."),
    ("h2", "Step 1 — Cliente"),
    ("bullets",
     "Nome (obbligatorio)",
     "Telefono",
     "Targa",
     "Tipo veicolo: Auto, SUV, Furgone, Camper, Truck, Trailer, Bus, Moto, Altro"),
    ("h2", "Step 2 — Pneumatici"),
    ("p", "Aggiungi i pneumatici con uno dei tre metodi:"),
    ("bullets",
     "🔍 Scanner QR/EAN → fotocamera. Riconosce: URL EPREL, EAN-13, DOT, "
     "JSON QuoteFlash, testo con misura",
     "＋ Aggiungi a mano → form con marca, misura ETRTO, quantità, "
     "prezzo, DOT",
     "Provenienza da analisi AI → automatico se sei arrivato da "
     "'Crea preventivo da questa analisi'"),
    ("p",
     "Tocca una card pneumatico per modificarne quantità o prezzo. "
     "Tocca la ✕ per rimuoverlo."),
    ("h2", "Fornitori B2B"),
    ("p",
     "I chip in fondo allo step 2 sono deep-link verso fornitori B2B. "
     "Toccando un chip si apre la ricerca del fornitore con la misura del "
     "primo pneumatico già compilata. Default: EPREL, Ciavarella B2B, "
     "Intergomma B2B, Google Shop. Puoi aggiungere fornitori personalizzati "
     "in Impostazioni."),
    ("h2", "Step 3 — Servizi"),
    ("p",
     "Catalogo a griglia. Tocca una tile per selezionarla. Le tiles in "
     "modalità AUTO includono Montaggio, Equilibratura, Convergenza, "
     "Assetto 3D, Gonfiaggio azoto, Stoccaggio stagionale, Smaltimento. "
     "In modalità TRUCK trovi Montaggio truck, Geometria assi, "
     "Riscolpitura, Soccorso stradale, ecc. Tutti i prezzi sono "
     "personalizzabili in Impostazioni."),
    ("h2", "Step 4 — Output"),
    ("p", "Riepilogo con totale. Quattro azioni:"),
    ("bullets",
     "📄 Scarica PDF — preventivo formattato pronto da stampare",
     "📋 Copia testo — copia il preventivo in formato testo per incollarlo",
     "💬 WhatsApp — apre WhatsApp con il preventivo già scritto, usando "
     "il numero del cliente (o il default da Impostazioni)",
     "📡 Webhook Make.com — invia i dati al webhook configurato"),

    ("h1", "7. Storico"),
    ("p",
     "Tutto quello che fai (analisi + preventivi) viene salvato automaticamente. "
     "La schermata Storico mostra tutto in ordine cronologico."),
    ("h2", "Filtri e ricerca"),
    ("bullets",
     "Tab 'Tutto / Analisi / Preventivi' per filtrare per tipo",
     "Campo di ricerca testuale (cerca per cliente, targa, misura, esito, ecc.)"),
    ("h2", "Dettaglio elemento"),
    ("p",
     "Toccando una voce si apre un bottom-sheet con tutti i dati. Da lì puoi:"),
    ("bullets",
     "Apri — riapre la schermata analisi (per riscaricare PDF, ecc.)",
     "Elimina — rimuove la voce dallo storico"),
    ("h2", "Export e pulizia"),
    ("bullets",
     "⬇ CSV — esporta tutto lo storico in formato CSV (Excel-friendly)",
     "⬇ JSON — esporta in JSON (backup completo)",
     "🗑 Pulisci — cancella TUTTO lo storico (con conferma)"),

    ("h1", "8. Officine vicine"),
    ("p",
     "Mappa di gommisti e officine intorno alla tua posizione. Usa OpenStreetMap "
     "via Overpass API: è gratuita e non richiede API key."),
    ("h2", "Come funziona"),
    ("bullets",
     "Al primo accesso chiede il permesso di geolocalizzazione",
     "Se neghi, puoi inserire manualmente una città o un CAP",
     "Filtra per categoria: Tutto, Gommisti, Officine, Auto",
     "Toccando un risultato si apre Google Maps con la posizione precisa"),

    ("h1", "9. Impostazioni"),
    ("p",
     "Nove sezioni a fisarmonica. Tocca l'intestazione di ogni sezione per "
     "espanderla/comprimerla."),
    ("h2", "1 · Profilo officina"),
    ("p", "Nome, P.IVA, telefono, indirizzo, email. Compaiono nelle intestazioni PDF."),
    ("h2", "2 · Provider AI"),
    ("p", "Cinque provider supportati con modelli ed endpoint configurabili:"),
    ("table",
     ["Provider", "Modelli", "Note"],
     ["Google Gemini", "gemini-2.5-flash-lite, gemini-2.5-flash", "Gratuita per uso personale"],
     ["OpenAI", "gpt-4o-mini, gpt-4o", "Richiede billing"],
     ["Anthropic Claude", "claude-haiku-4-5, claude-sonnet-4-6", "Browser direct-access"],
     ["Mistral", "pixtral-12b-2409, pixtral-large-latest", "Visione multimodale"],
     ["Ollama locale", "llava, gemma3, minicpm-v", "Offline, gira sul tuo computer"]),
    ("p",
     "Il pulsante 'Testa connessione' chiama il provider e verifica che la "
     "chiave funzioni. Endpoint custom = puoi sovrascrivere l'URL di default "
     "(utile per proxy aziendali o gateway interni)."),
    ("h2", "3 · EPREL Proxy"),
    ("p",
     "URL di un Cloudflare Worker che recupera i dati EPREL completi da un "
     "QR code EPREL. Lascia vuoto se non lo hai: l'app aprirà direttamente "
     "il sito EPREL UE in una nuova scheda."),
    ("h2", "4 · Webhook Make.com"),
    ("p",
     "URL del webhook Make.com (Integromat) per inviare automaticamente "
     "analisi e preventivi a sistemi esterni (CRM, fogli Google, ecc.)."),
    ("h2", "5 · WhatsApp"),
    ("p",
     "Numero di default da usare per l'invio WhatsApp quando il cliente non "
     "ne ha uno proprio."),
    ("h2", "6 · Fornitori B2B"),
    ("p",
     "Lista dei deep-link verso portali B2B. Ogni fornitore ha un URL template "
     "con segnaposto {w}, {h}, {r} che vengono sostituiti con la misura del "
     "pneumatico. Aggiungi/rimuovi con i pulsanti dedicati. Pulsante 'Ripristina "
     "default' rimette EPREL, Ciavarella, Intergomma e Google Shop."),
    ("h2", "7 · Modalità default"),
    ("p", "Sceglie se l'app si apre in modalità AUTO o TRUCK."),
    ("h2", "8 · Avanzate"),
    ("bullets",
     "⬇ Esporta backup JSON — scarica tutti i dati locali",
     "⬆ Importa backup JSON — ripristina da un backup precedente",
     "🗑 Reset app — cancella tutto e ricarica (con conferma)"),
    ("h2", "9 · Disclaimer e note legali"),
    ("p",
     "Testo completo del disclaimer open source (13 sezioni: natura del progetto e "
     "rilascio MIT, valore indicativo della diagnosi AI, esclusione dall'AI Act "
     "ad alto rischio, responsabilità sui dati inseriti, costi di terze parti, "
     "servizi esterni, conservazione dati, privacy/GDPR/localStorage, marchi "
     "di terzi, sicurezza informatica, limitazione di responsabilità, evoluzione "
     "del progetto, legge italiana e foro). Sempre consultabile da qui. Il "
     "pulsante 'Riapri schermata di presa visione' permette di rileggere il "
     "disclaimer con la checkbox originale."),
    ("h2", "10 · Info"),
    ("p", "Versione corrente e link al sito PezzaliApp."),

    ("h1", "10. Installare come PWA"),
    ("p", "TireCheckTire è installabile sulla Home del tuo telefono come una normale app."),
    ("h2", "iPhone / iPad (Safari)"),
    ("bullets",
     "Apri il sito in Safari",
     "Tocca l'icona Condividi (quadratino con freccia)",
     "Scorri e tocca 'Aggiungi a Home'",
     "Conferma con 'Aggiungi'"),
    ("h2", "Android (Chrome)"),
    ("bullets",
     "Apri il sito in Chrome",
     "Tocca il menu ⋮ in alto a destra",
     "Tocca 'Installa app' o 'Aggiungi alla schermata Home'",
     "Conferma"),
    ("h2", "Desktop (Chrome, Edge)"),
    ("bullets",
     "Apri il sito",
     "Nella barra degli indirizzi compare un'icona di installazione",
     "Clicca → 'Installa'"),
    ("callout",
     "Una volta installata, l'app funziona offline per le funzioni che "
     "non richiedono Internet (storico, generazione PDF, inserimento dati). "
     "Per l'AI, le officine vicine e i webhook serve Internet."),

    ("h1", "11. Disclaimer e note legali"),
    ("p",
     "TireCheckTire è un progetto open source rilasciato gratuitamente a "
     "titolo personale da Alessandro Pezzali, persona fisica, al di fuori "
     "di qualsiasi attività professionale, commerciale o d'impresa. Il "
     "codice è su github.com/pezzaliapp/TireCheckTire sotto licenza MIT. "
     "Non sussiste alcun rapporto contrattuale tra l'autore e gli utenti "
     "dell'applicazione."),
    ("p",
     "All'apertura dell'app — e ogni volta che la versione del documento "
     "cambia — compare una schermata bloccante che chiede di prendere "
     "visione del disclaimer. L'app non funziona finché la casella non è "
     "spuntata e il pulsante 'Accetta e continua' premuto."),
    ("p", "Il disclaimer copre in particolare:"),
    ("bullets",
     "Natura del progetto: rilascio personale gratuito sotto licenza MIT, "
     "fornito \"AS IS\", senza garanzie",
     "Diagnosi AI: valore INDICATIVO, da verificare sempre manualmente",
     "AI Act (Reg. UE 2024/1689): l'app NON è un sistema AI ad alto "
     "rischio e non prende decisioni autonome di sicurezza",
     "Inserimento dati: responsabilità esclusiva dell'utente",
     "Costi di terze parti (Gemini, OpenAI, ecc.): a carico dell'utente",
     "Servizi esterni: disponibilità e continuità non garantite",
     "Conservazione dati: backup a carico dell'utente (solo localStorage, "
     "nessun server proprietario)",
     "Privacy/GDPR: nessun cookie di tracciamento; l'utente che tratta dati "
     "di terzi è Titolare del Trattamento",
     "Marchi di terzi: uso descrittivo ex art. 21 CPI (D.Lgs. 30/2005)",
     "Sicurezza informatica: responsabilità sul dispositivo dell'utente",
     "Limitazione di responsabilità: come da legge italiana, fatta salva "
     "responsabilità per dolo o colpa grave ex art. 1229 c.c.",
     "Foro: luogo di residenza dell'autore, fatto salvo foro del consumatore"),
    ("p",
     "Il testo integrale è sempre consultabile da Impostazioni → "
     "'9 · Disclaimer e note legali'. Una versione sintetica del disclaimer "
     "viene stampata anche in fondo a ogni PDF (report di analisi e "
     "preventivi) come tutela aggiuntiva."),
    ("callout",
     "L'app non sostituisce mai l'ispezione tecnica dell'operatore. La "
     "diagnosi AI è un ausilio, non una perizia: prima di sostituire "
     "un pneumatico o rilasciare un documento, verifica sempre i dati."),

    ("h1", "12. Privacy"),
    ("bullets",
     "Nessun tracking, nessuna analytics di terze parti",
     "Foto e dati cliente restano sul tuo dispositivo (localStorage)",
     "Le API key dei provider AI sono salvate SOLO nel tuo browser",
     "Le chiamate AI vanno SOLO all'endpoint che hai configurato",
     "PezzaliApp non riceve alcun dato"),

    ("h1", "13. Risoluzione problemi"),
    ("table",
     ["Sintomo", "Possibile causa", "Soluzione"],
     ["Pulsanti in basso non appaiono",
      "Service Worker che serve una versione vecchia",
      "Pulisci cache del sito (Impostazioni Safari → Dati siti web) e ricarica"],
     ["La fotocamera non si apre al tap",
      "Permesso negato o browser obsoleto",
      "Verifica i permessi del sito; usa Safari/Chrome aggiornati"],
     ["L'AI risponde sempre con errore",
      "API key mancante o errata",
      "Impostazioni → Provider AI → reinserisci la chiave e premi 'Testa'"],
     ["L'OCR del fianco non trova niente",
      "Foto sfocata o angolata",
      "Avvicinati al fianco, luce uniforme, mantieni la foto frontale"],
     ["Officine vicine: lista vuota",
      "Geolocalizzazione negata",
      "Usa 'Inserisci città' per ricerca manuale"],
     ["PDF non si scarica",
      "Mancanza firma e modo 'senza firma' non attivo",
      "Firma il documento o premi 'Procedi senza firma'"]),

    ("h1", "14. Crediti"),
    ("p",
     "TireCheckTire — The Tire Intelligence Suite è sviluppato da "
     "PezzaliApp. Versione 1.0, maggio 2026."),
    ("p", "Licenza: MIT. Codice sorgente: github.com/pezzaliapp/TireCheckTire"),
    ("p", "Suite nata dall'unione di tre PWA: QuoteFlash, TireCheck Pro, TireCheck Fleet."),
]


# ---------------------------------------------------------------------------
# DOCX renderer
# ---------------------------------------------------------------------------

ACCENT_HEX = 0xE8FF47
DARK_HEX   = 0x0A0A0A
GRAY_HEX   = 0x666666


def build_docx(out_path: Path):
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Helvetica'
    style.font.size = Pt(11)

    for item in CONTENT:
        kind = item[0]

        if kind == 'cover':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('🛞')
            run.font.size = Pt(72)
            doc.add_paragraph()

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(DOC_TITLE)
            run.bold = True
            run.font.size = Pt(28)

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(DOC_SUBTITLE)
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            for _ in range(2):
                doc.add_paragraph()

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(DOC_BYLINE)
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            doc.add_page_break()

        elif kind == 'h1':
            heading = doc.add_heading(item[1], level=1)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
                run.font.size = Pt(20)

        elif kind == 'h2':
            heading = doc.add_heading(item[1], level=2)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
                run.font.size = Pt(14)

        elif kind == 'p':
            p = doc.add_paragraph(item[1])
            p.paragraph_format.space_after = Pt(8)

        elif kind == 'bullets':
            for bullet in item[1:]:
                doc.add_paragraph(bullet, style='List Bullet')

        elif kind == 'callout':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run('💡  ' + item[1])
            run.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)

        elif kind == 'table':
            header_row = item[1]
            body_rows = item[2:]
            table = doc.add_table(rows=1 + len(body_rows), cols=len(header_row))
            table.style = 'Light Grid Accent 1'
            for col_idx, header in enumerate(header_row):
                cell = table.rows[0].cells[col_idx]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for row_idx, row_data in enumerate(body_rows, start=1):
                for col_idx, value in enumerate(row_data):
                    table.rows[row_idx].cells[col_idx].text = str(value)
            doc.add_paragraph()

    # Footer (in-document, not page footer)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('—  fine  —')
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.italic = True

    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# PDF renderer
# ---------------------------------------------------------------------------

# Register a Unicode-rich font for the PDF so emojis and symbols render
# instead of falling back to "missing glyph" boxes. Tries macOS Arial Unicode
# first (covers virtually everything), then DejaVu, then gives up gracefully.
PDF_FONT = 'Helvetica'  # safe default
PDF_FONT_BOLD = 'Helvetica-Bold'

_CANDIDATES = [
    ('/Library/Fonts/Arial Unicode.ttf', 'ArialUnicode'),
    ('/System/Library/Fonts/Supplemental/Arial Unicode.ttf', 'ArialUnicode'),
    ('/Library/Fonts/Arial.ttf', 'ArialUnicode'),
    ('/System/Library/Fonts/Supplemental/Arial.ttf', 'ArialUnicode'),
    ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 'DejaVu'),
    ('/Library/Fonts/DejaVuSans.ttf', 'DejaVu'),
]

for path, name in _CANDIDATES:
    if Path(path).exists():
        try:
            pdfmetrics.registerFont(TTFont(name, path))
            PDF_FONT = name
            # Try the bold variant in the same directory; otherwise reuse the regular.
            bold_path = path.replace('.ttf', ' Bold.ttf').replace('Arial Unicode', 'Arial Bold')
            bold_alt  = Path(path).with_name(Path(path).stem + '-Bold.ttf')
            if Path(bold_path).exists():
                pdfmetrics.registerFont(TTFont(name + '-Bold', bold_path))
                PDF_FONT_BOLD = name + '-Bold'
            elif bold_alt.exists():
                pdfmetrics.registerFont(TTFont(name + '-Bold', str(bold_alt)))
                PDF_FONT_BOLD = name + '-Bold'
            else:
                PDF_FONT_BOLD = name  # synthetic bold via reportlab
            break
        except Exception as exc:
            print(f'! font {name} skipped: {exc}')

print(f'PDF font: {PDF_FONT} (bold: {PDF_FONT_BOLD})')


# Even with Arial Unicode, the Supplementary Multilingual Plane emojis
# (camera, clipboard, magnifier, etc. — all U+1F3xx and above) are not
# included. Map them to BMP glyphs the font DOES carry.
PDF_GLYPH = {
    '🛞': '',     # cover image is used instead
    '📸': '▸',
    '📷': '▸',
    '📋': '▸',
    '🔍': '▸',
    '🔎': '▸',
    '📍': '▸',
    '💡': '✦',
    '✅': '✓',
    '⚠':  '!',
    '🔴': '●',
    '🤖': '▸',
    '📄': '▸',
    '📡': '▸',
    '⬇':  '↓',
    '⬆':  '↑',
    '🗑': '×',
    '🔧': '▸',
    '⚙':  '*',
    '⚙️': '*',
    '🏠': '▸',
    '🗂':  '▸',
    '⚖':  '▸',
    '⚖️': '▸',
    '🎯': '▸',
    '📐': '▸',
    '💨': '▸',
    '🔩': '▸',
    '📦': '▸',
    '🔵': '●',
    '♻':  '↻',
    '♻️': '↻',
    '🪓': '▸',
    '🚛': '►',
    '🚗': '►',
    '💬': '▸',
    '✕':  '×',
    '＋': '+',
    '🇪🇺': 'UE',
    '🌙': '☾',
    '☀':  '☼',
    '🌆': '☼',
    '👋': '',
}


def pdf_safe(text: str) -> str:
    for emoji, replacement in PDF_GLYPH.items():
        text = text.replace(emoji, replacement)
    return ' '.join(text.split())


def build_pdf(out_path: Path):
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=22 * mm,
        title=DOC_TITLE,
        author=DOC_BYLINE,
    )

    styles = getSampleStyleSheet()
    accent = HexColor(f'#{ACCENT_HEX:06x}')
    dark   = HexColor(f'#{DARK_HEX:06x}')
    gray   = HexColor(f'#{GRAY_HEX:06x}')

    title_style = ParagraphStyle(
        'Title', parent=styles['Title'],
        fontName=PDF_FONT_BOLD, fontSize=28, leading=34, alignment=TA_CENTER,
        textColor=dark, spaceAfter=4)
    subtitle_style = ParagraphStyle(
        'Subtitle', parent=styles['Normal'],
        fontName=PDF_FONT, fontSize=13, leading=16, alignment=TA_CENTER,
        textColor=gray, spaceAfter=18)
    byline_style = ParagraphStyle(
        'Byline', parent=styles['Italic'],
        fontName=PDF_FONT, fontSize=11, alignment=TA_CENTER, textColor=gray, spaceAfter=4)
    h1_style = ParagraphStyle(
        'H1', parent=styles['Heading1'],
        fontName=PDF_FONT_BOLD, fontSize=20, leading=24, textColor=dark,
        spaceBefore=18, spaceAfter=10, keepWithNext=True)
    h2_style = ParagraphStyle(
        'H2', parent=styles['Heading2'],
        fontName=PDF_FONT_BOLD, fontSize=14, leading=18, textColor=dark,
        spaceBefore=12, spaceAfter=6, keepWithNext=True)
    body_style = ParagraphStyle(
        'Body', parent=styles['BodyText'],
        fontName=PDF_FONT, fontSize=11, leading=15, textColor=dark,
        spaceAfter=8, alignment=TA_LEFT)
    bullet_style = ParagraphStyle(
        'Bullet', parent=body_style, fontName=PDF_FONT,
        leftIndent=18, bulletIndent=6, spaceAfter=4)
    callout_style = ParagraphStyle(
        'Callout', parent=body_style, fontName=PDF_FONT,
        backColor=HexColor('#fffae0'),
        borderColor=accent, borderWidth=0.5, borderPadding=8,
        leftIndent=0, rightIndent=0, fontSize=10.5, leading=15,
        spaceBefore=6, spaceAfter=10)
    flow = []

    logo_path = Path(__file__).parent.parent / 'assets' / 'icon-192.png'

    for item in CONTENT:
        kind = item[0]

        if kind == 'cover':
            flow.append(Spacer(1, 40 * mm))
            if logo_path.exists():
                img = Image(str(logo_path), width=42 * mm, height=42 * mm)
                img.hAlign = 'CENTER'
                flow.append(img)
                flow.append(Spacer(1, 12 * mm))
            flow.append(Paragraph(DOC_TITLE, title_style))
            flow.append(Paragraph(DOC_SUBTITLE, subtitle_style))
            flow.append(Spacer(1, 24 * mm))
            flow.append(Paragraph(DOC_BYLINE, byline_style))
            flow.append(PageBreak())

        elif kind == 'h1':
            flow.append(Paragraph(pdf_safe(item[1]), h1_style))

        elif kind == 'h2':
            flow.append(Paragraph(pdf_safe(item[1]), h2_style))

        elif kind == 'p':
            flow.append(Paragraph(pdf_safe(item[1]), body_style))

        elif kind == 'bullets':
            list_items = [ListItem(Paragraph(pdf_safe(b), bullet_style), leftIndent=12) for b in item[1:]]
            flow.append(ListFlowable(list_items, bulletType='bullet',
                                     start='•', leftIndent=12, bulletFontSize=10))
            flow.append(Spacer(1, 6))

        elif kind == 'callout':
            flow.append(Paragraph('✦ ' + pdf_safe(item[1]), callout_style))

        elif kind == 'table':
            # Wrap each cell in a Paragraph so text wraps within column width
            # (raw strings would overflow into the neighbouring column).
            cell_style = ParagraphStyle(
                'TableCell', parent=body_style, fontName=PDF_FONT,
                fontSize=9.5, leading=12, spaceAfter=0, spaceBefore=0)
            header_style = ParagraphStyle(
                'TableHead', parent=body_style, fontName=PDF_FONT_BOLD,
                fontSize=9.5, leading=12, textColor=accent,
                spaceAfter=0, spaceBefore=0)
            header_row = [Paragraph(pdf_safe(c), header_style) for c in item[1]]
            body_rows = [[Paragraph(pdf_safe(c), cell_style) for c in row] for row in item[2:]]
            data = [header_row] + body_rows

            avail = A4[0] - 40 * mm
            n_cols = len(item[1])
            # Bias toward wider rightmost column for "description"-style tables
            if n_cols == 3:
                col_w = [avail * 0.22, avail * 0.30, avail * 0.48]
            elif n_cols == 2:
                col_w = [avail * 0.18, avail * 0.82]
            else:
                col_w = [avail / n_cols] * n_cols

            tbl = Table(data, colWidths=col_w, repeatRows=1)
            tbl.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), dark),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [HexColor('#fafafa'), HexColor('#f0f0f0')]),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 6),
                ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('LINEBELOW', (0, 0), (-1, 0), 0.4, dark),
                ('GRID', (0, 1), (-1, -1), 0.25, HexColor('#dddddd')),
            ]))
            flow.append(tbl)
            flow.append(Spacer(1, 10))

    # Page footer with brand
    def footer(canvas, doc_):
        canvas.saveState()
        canvas.setFont(PDF_FONT, 8)
        canvas.setFillColor(gray)
        canvas.drawString(20 * mm, 12 * mm, 'TireCheckTire · Manuale d\'uso · v1.0')
        canvas.drawRightString(A4[0] - 20 * mm, 12 * mm, f'Pag. {doc_.page}')
        canvas.restoreState()

    doc.build(flow, onFirstPage=footer, onLaterPages=footer)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    out_dir = Path(__file__).parent
    docx_path = out_dir / 'Manuale-TireCheckTire.docx'
    pdf_path  = out_dir / 'Manuale-TireCheckTire.pdf'

    build_docx(docx_path)
    print(f'✓ DOCX: {docx_path}  ({docx_path.stat().st_size // 1024} KB)')

    build_pdf(pdf_path)
    print(f'✓ PDF:  {pdf_path}  ({pdf_path.stat().st_size // 1024} KB)')


if __name__ == '__main__':
    main()
